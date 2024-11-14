import os
import random
import uuid
import tempfile
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from faker import Faker
import math2docx
import matplotlib.pyplot as plt

# настройка Faker для генерации фиктивного текста на русском языке
fake = Faker('ru_RU')

# константы для цвета фона таблиц
colors_list = [
    # базовые цвета
    'FFFFFF',  # White
    '000000',  # Black
        
    # оттенки серого
    'F5F5F5',  # WhiteSmoke
    'E0E0E0',  # Light Gray
    'CCCCCC',  # Silver
    '808080',  # Gray
        
    # синие тона
    'ECF0F1',  # Cloud
    'D6EAF8',  # Light Blue
    'AED6F1',  # Soft Blue
    '3498DB',  # Dodger Blue
    '2E86C1',  # Ocean Blue
        
    # зеленые тона
    'E8F5E9',  # Mint Cream
    'A2D9CE',  # Light Sea Green
    '27AE60',  # Nephritis
    '1E8449',  # Forest Green
       
    # красные тона
    'FADBD8',  # Misty Rose
    'F1948A',  # Light Coral
    'E74C3C',  # Alizarin
    'C0392B',  # Dark Red
        
    # желтые и оранжевые тона
    'FEF9E7',  # Light Yellow
    'FCF3CF',  # Cream
    'F7DC6F',  # Sandstone
    'F39C12',  # Orange
        
    # фиолетовые тона
    'F4ECF7',  # Lavender Mist
    'D7BDE2',  # Plum
    '8E44AD',  # Purple
        
    # коричневые тона
    'FDEBD0',  # Antique White
    'E59866',  # Peru
    'BA4A00'   # Saddle Brown
]

class DocumentState:
    """Класс для отслеживания состояния документа, включая текущие колонки, ориентацию и т.д."""
    def __init__(self, document):
        self.current_columns = 1  # текущие колонки (1, 2 или 3)
        self.document = document
        self.multicol_sections_added = {'2': False, '3': False}
        self.elements_in_multicol = 0  # счётчик элементов в многоколонной секции
        self.max_elements_multicol = 10  # максимальное количество элементов в секции
        self.min_elements_per_column = 2  # минимальное количество элементов на колонку
        self.required_elements = 0  # общее требуемое количество элементов
        self.in_multicol = False  # флаг, находится ли документ сейчас в многоколонной секции
        self.current_orientation = WD_ORIENT.PORTRAIT  # текущая ориентация
        self.multicol_table = None  # таблица, используемая для многоколонной секции
        self.current_multicol_cell = None  # текущая ячейка таблицы для вставки контента
        self.multicol_col_index = 0  # индекс текущей колонки
        self.column_elements = []

    def can_add_multicolumn(self, num_cols):
        """Проверяет, можно ли добавить секцию с num_cols колонками."""
        return not self.multicol_sections_added.get(str(num_cols), False)

    def start_multicolumn(self, num_cols):
        """Запускает многоколонную секцию с использованием полностью прозрачной таблицы."""
        if not self.can_add_multicolumn(num_cols):
            return False

        self.document.add_paragraph()

        table = self.document.add_table(rows=1, cols=num_cols)
        table.autofit = False

        # используем практически всю ширину страницы
        section = self.document.sections[-1]
        page_width = section.page_width
        margin_left = Inches(1.2)  
        margin_right = Inches(0.6)  
        
        # используем максимальную возможную ширину
        usable_width = page_width - margin_left - margin_right
        column_width = usable_width / num_cols

        for row in table.rows:
            for cell in row.cells:
                cell.width = column_width
                # убираем все отступы внутри ячеек
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = 0
                    paragraph.paragraph_format.space_before = 0
                    paragraph.paragraph_format.left_indent = 0
                    paragraph.paragraph_format.right_indent = 0

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        self.multicol_table = table
        self.current_multicol_cell = table.rows[0].cells[0]
        self.multicol_col_index = 0
        self.current_columns = num_cols
        self.in_multicol = True
        self.elements_in_multicol = 0
        self.multicol_sections_added[str(num_cols)] = True
        self.column_elements = [0] * num_cols

        return True

    def end_multicolumn(self):
        """Заканчивает текущую многоколонную секцию."""
        if not self.in_multicol:
            return

        # добавляем новую секцию
        new_section = self.document.add_section(WD_SECTION_START.CONTINUOUS)
        
        # добавляем пустой параграф в новой секции
        p = self.document.add_paragraph()
        
        # сбрасываем состояния
        self.multicol_table = None
        self.current_multicol_cell = None
        self.multicol_col_index = 0
        self.current_columns = 1
        self.in_multicol = False
        self.elements_in_multicol = 0
        self.required_elements = 0
        self.column_elements = []

    def get_current_multicol_cell(self):
        """Возвращает текущую ячейку таблицы для вставки контента."""
        if not self.in_multicol or not self.multicol_table:
            return None
            
        # проверяем, есть ли колонки без элементов
        empty_columns = [i for i, count in enumerate(self.column_elements) if count == 0]
        
        if empty_columns:
            # если есть пустые колонки, берем первую из них
            self.multicol_col_index = empty_columns[0]
        else:
            # если все колонки имеют элементы, используем обычную ротацию
            self.multicol_col_index = (self.multicol_col_index + 1) % self.current_columns
            
        cell = self.multicol_table.rows[0].cells[self.multicol_col_index]
        self.column_elements[self.multicol_col_index] += 1
        return cell

    def is_column_filled(self):
        """Проверяет, заполнена ли текущая колонка минимальным количеством элементов"""
        elements_in_current_column = self.elements_in_multicol % self.current_columns
        return elements_in_current_column >= self.min_elements_per_column

    def can_end_multicolumn(self):
        """Проверяет, можно ли завершить многоколонную секцию"""
        # проверяем, что все колонки содержат хотя бы один элемент
        has_empty_columns = any(count == 0 for count in self.column_elements)
        min_elements_satisfied = all(count >= self.min_elements_per_column for count in self.column_elements)
        total_elements_satisfied = self.elements_in_multicol >= self.required_elements
        
        return not has_empty_columns and min_elements_satisfied and total_elements_satisfied

    def can_start_multicolumn(self):
        """Проверяет, можно ли начать новую многоколонную секцию."""
        return (self.can_add_multicolumn(2) or self.can_add_multicolumn(3))

def add_table_caption(document, before=True):
    """Добавляет подпись к таблице перед или после таблицы."""
    caption_type = random.choice(['Табл. {num}. {text}', 'Таблица {num} - {text}', 'Таблица. {text}'])
    table_num = random.randint(1, 100)
    if 'num' in caption_type:
        if caption_type.startswith('Таблица.'):
            caption_text = caption_type.format(text=fake.sentence(nb_words=4))
        else:
            caption_text = caption_type.format(num=table_num, text=fake.sentence(nb_words=4))
    else:
        caption_text = f"Таблица. {fake.sentence(nb_words=4)}"

    p = document.add_paragraph(caption_text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # настройка цвета шрифта (синий или чёрный)
    color = random.choice(['blue', 'black'])
    if color == 'blue':
        p.runs[0].font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1  # стандартный голубоватый
    else:
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # чёрный

    # настройка размера шрифта
    p.runs[0].font.size = Pt(random.randint(10, 12))

    # настройка начертания (курсивное или прямое)
    p.runs[0].font.italic = random.choice([True, False])

    # настройка жирности шрифта (жирный или простой)
    p.runs[0].font.bold = random.choice([True, False])

def add_table(document, doc_state):
    """Добавляет таблицу с подписью в документ, случайно располагая подпись перед или после таблицы."""
    if doc_state.in_multicol:
        # в многоколонной секции таблицы добавлять нельзя
        return

    # решаем, где разместить подпись: перед или после
    caption_before = random.choice([True, False])

    # добавление подписи
    if caption_before:
        add_table_caption(document, before=True)

    # добавление самой таблицы
    table_styles = get_table_styles(document)
    num_rows = random.randint(5, 10)
    num_cols = random.randint(3, 6)

    # создаём таблицу
    table = document.add_table(rows=num_rows, cols=num_cols)
    
    # выбор стиля таблицы согласно заданным вероятностям
    style_choice = random.random() * 100  # Генерируем число от 0 до 100

    if style_choice < 45:  # 45% вероятность белой таблицы
        table_style_option = 'white'
        # для белых таблиц всегда используем стиль с границами
        table.style = 'Table Grid'
    else:
        # для остальных типов таблиц используем случайный стиль
        table.style = random.choice(table_styles) if table_styles else 'Table Grid'
        
        if style_choice < 70:  # 25% вероятность одноцветной таблицы
            table_style_option = 'single_color'
        elif style_choice < 95:  # 25% вероятность чередующихся цветов
            table_style_option = 'alternating_color'
        else:  # 5% вероятность случайных цветов
            table_style_option = 'random'

    table.alignment = random.choice([
        WD_TABLE_ALIGNMENT.LEFT,
        WD_TABLE_ALIGNMENT.CENTER,
        WD_TABLE_ALIGNMENT.RIGHT
    ])
    table.autofit = False

    # установка ширины таблицы
    table_width = Cm(18)

    # инициализация цветов для alternating_color
    if table_style_option == 'alternating_color':
        # выбираем два разных цвета, исключая белый (FFFFFF)
        available_colors = [c for c in colors_list if c != 'FFFFFF']
        color_a, color_b = random.sample(available_colors, 2)
    elif table_style_option == 'single_color':
        # выбираем один цвет, исключая белый (FFFFFF)
        available_colors = [c for c in colors_list if c != 'FFFFFF']
        single_color = random.choice(available_colors)

    for row_idx, row in enumerate(table.rows):
        row.height = Pt(12)
        for idx, cell in enumerate(row.cells):
            cell.width = table_width / num_cols
            cell.text = ""

            # Применяем выбранный стиль
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')

            if table_style_option == 'white':
                shading_elm.set(qn('w:fill'), 'FFFFFF')
            elif table_style_option == 'single_color':
                shading_elm.set(qn('w:fill'), single_color)
            elif table_style_option == 'alternating_color':
                row_color = color_a if row_idx % 2 == 0 else color_b
                shading_elm.set(qn('w:fill'), row_color)
            else:  # random
                fill_color = random.choice([c for c in colors_list if c != 'FFFFFF'])
                shading_elm.set(qn('w:fill'), fill_color)

            cell._tc.get_or_add_tcPr().append(shading_elm)

            # добавляем случайный текст
            cell.text = fake.text(max_nb_chars=random.randint(5, 50)) if random.choice([True, False]) else str(fake.random_number(digits=5))

            # настройка выравнивания текста в ячейках
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.alignment = random.choice([
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.RIGHT,
                WD_ALIGN_PARAGRAPH.JUSTIFY
            ])

            # настройка шрифта
            if cell_paragraph.runs:
                run = cell_paragraph.runs[0]
                run.font.size = Pt(random.randint(8, 12))

    table._element.getparent().spacing_before = Pt(random.randint(4, 12))
    table._element.getparent().spacing_after = Pt(random.randint(4, 12))

    if not caption_before:
        add_table_caption(document, before=False)

def add_footnote(paragraph, text, footnote_number):
    """Добавляет сноску в абзац."""
    run = paragraph.add_run(f"[{footnote_number}]")
    run.font.superscript = True
    run.font.size = Pt(8)
    paragraph.add_run(f" {text}")

def add_document_end_footnotes(document, end_footnotes):
    """Добавляет список концевых сносок в конце документа на новой странице с портретной ориентацией."""
    if not end_footnotes:
        return

    # добавляем пустой параграф перед новой секцией для гарантии разрыва
    last_p = document.add_paragraph()
    last_p.runs.clear()  
    
    # добавляем явный разрыв страницы
    run = last_p.add_run()
    run.add_break(WD_BREAK.PAGE)

    # добавляем раздел перед концевыми сносками
    new_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    
    # разрываем связь с предыдущей секцией
    new_section.start_type = WD_SECTION_START.NEW_PAGE
    new_section.link_to_previous = False

    # сброс параметров страницы к стандартным значениям (A4 портретная)
    new_section.page_width = Mm(210)  # ширина A4
    new_section.page_height = Mm(297)  # высота A4
    new_section.orientation = WD_ORIENT.PORTRAIT

    # устанавливаем отступы страницы
    new_section.left_margin = Mm(25.4)
    new_section.right_margin = Mm(25.4)
    new_section.top_margin = Mm(25.4)
    new_section.bottom_margin = Mm(25.4)

    # дополнительно вызываем функцию установки ориентации
    set_section_orientation(new_section, WD_ORIENT.PORTRAIT)

    # добавляем заголовок "Концевые сноски"
    p = document.add_paragraph("Концевые сноски", style='Heading 2')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    
    if p.runs:
        p.runs[0].font.size = Pt(random.randint(14, 18))
        p.runs[0].font.bold = True
        
        # случайный выбор цвета: голубой Word (Accent 1) или черный
        if random.choice([True, False]):
            p.runs[0].font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1
        else:
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # добавляем сами сноски
    for idx, footnote in enumerate(end_footnotes, 1):
        p = document.add_paragraph(f"{idx}. {footnote}", style='Normal')
        p.paragraph_format.space_before = Pt(random.randint(0, 2))
        p.paragraph_format.space_after = Pt(random.randint(0, 2))
        if p.runs:
            p.runs[0].font.size = Pt(random.randint(10, 14))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_image(document, doc_state):
    """Добавляет изображение с подписью в документ, с возможностью изменения размера изображения и различными стилями подписи."""
    pic_path = generate_random_plot_image(doc_state)
    if not os.path.exists(pic_path):
        return

    img_width = min(Inches(random.uniform(4, 6)), Inches(6))  # максимальная ширина 6 дюймов

    # случайное определение текста подписи
    caption_variants = [
        f"Рис. {random.randint(1, 100)}",
        f"Рисунок {random.randint(1, 100)} -",
        f"Рисунок {random.randint(1, 100)}"
    ]
    caption_text = random.choice(caption_variants) + f" {fake.sentence(nb_words=5)}"

    # случайное определение уменьшения размера шрифта подписи
    caption_font_size = Pt(random.randint(8, 10))

    # случайное выравнивание подписи (чаще по центру)
    caption_alignment = random.choices(
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        weights=[0.6, 0.2, 0.2]
    )[0]

    # функция добавления изображения с уменьшением размера при необходимости
    def add_resizable_image(container, width):
        run = container.add_run()
        try:
            run.add_picture(pic_path, width=width)
            return True
        except Exception as e:
            return False

    # чаще добавляем подпись после изображения
    caption_before = random.choices([True, False], weights=[0.4, 0.6])[0]

    # создание контейнера для изображения и подписи
    if doc_state.in_multicol:
        current_cell = doc_state.get_current_multicol_cell()
        if current_cell is None:
            return
        p_container = current_cell.add_paragraph()
        p_container.paragraph_format.keep_with_next = True
    else:
        p_container = document.add_paragraph()
        p_container.paragraph_format.keep_with_next = True

    # добавление подписи перед изображением, если выбрана эта опция
    if caption_before:
        p_caption = p_container.add_run(caption_text)
        p_caption.font.size = caption_font_size
        p_container.paragraph_format.keep_with_next = True
        p_container.alignment = caption_alignment
        p_container.add_run("\n")  # Добавляем перенос строки

    # добавление изображения и уменьшение его размера при необходимости
    while not add_resizable_image(p_container, img_width) and img_width > Cm(1):
        img_width -= Cm(0.5)

    # добавление подписи после изображения, если не выбрана опция "перед"
    if not caption_before:
        p_caption = p_container.add_run(f"\n{caption_text}")
        p_caption.font.size = caption_font_size
        p_container.paragraph_format.keep_with_next = True
        p_container.alignment = caption_alignment

    # удаление временного файла изображения
    os.remove(pic_path)

    # увеличение счётчика элементов в многоколонной секции
    if doc_state.in_multicol:
        doc_state.elements_in_multicol += 1
        if doc_state.elements_in_multicol >= doc_state.max_elements_multicol:
            doc_state.end_multicolumn()

def make_paragraph_keep_together(paragraph):
    """Применяет свойства для предотвращения разрыва страницы внутри абзаца."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    keep_together = OxmlElement('w:keepTogether')
    keep_together.set(qn('w:val'), '1')
    pPr.append(keep_together)

    keep_with_next = OxmlElement('w:keepNext')
    keep_with_next.set(qn('w:val'), '1')
    pPr.append(keep_with_next)

def make_run_keep_with_next(run):
    """Применяет свойства к запуску для предотвращения разрыва страницы между элементами."""
    r = run._r
    rPr = r.get_or_add_rPr()
    keep_with_next = OxmlElement('w:keepNext')
    keep_with_next.set(qn('w:val'), '1')
    rPr.append(keep_with_next)

def add_heading(document, level=2):
    """Добавляет заголовок заданного уровня в документ."""
    heading_text = fake.sentence(nb_words=6)
    p = document.add_heading(level=level)
    run = p.add_run(heading_text)
    run.font.size = Pt(random.randint(14 + (level - 1) * 2, 20 + (level - 1) * 2))
    
    # выбор цвета для текста: стандартный голубоватый или черный
    color_choice = random.choice(['standard_blue', 'almost_black'])
    if color_choice == 'standard_blue':
        run.font.color.rgb = RGBColor(0, 112, 192)  # стандартный голубоватый цвет
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)  # черный 

    if random.choice([True, False]):
        run.bold = True
    if random.choice([True, False]):
        run.italic = True
    
    heading_alignment = random.choice([
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT
    ])
    p.paragraph_format.space_before = Pt(12)
    p.alignment = heading_alignment

    make_paragraph_keep_together(p)
    make_run_keep_with_next(run)

def add_text(document, doc_state, text, font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Добавляет абзац текста в документ с заданным форматированием."""
    p = document.add_paragraph(text)
    run = p.runs[0]
    run.font.size = Pt(font_size)
    p.alignment = alignment

    # 50% шанс добавить красную строку
    if random.choice([True, False]):
        p.paragraph_format.first_line_indent = Cm(1.25)  
    else:
        p.paragraph_format.first_line_indent = Cm(0) 

    # разнообразие межстрочного интервала
    p.paragraph_format.space_after = Pt(random.randint(6, 12))

    p.space_after = Pt(random.randint(4, 12))
    p.space_before = Pt(random.randint(0, 2))

    # предотвращаем разрыв страницы внутри абзаца
    make_paragraph_keep_together(p)
    make_run_keep_with_next(run)

def add_text_to_cell(cell, text, font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Добавляет абзац текста в указанную ячейку таблицы с заданным форматированием."""
    p = cell.add_paragraph(text)
    run = p.runs[0]
    run.font.size = Pt(font_size)
    p.alignment = alignment

    # 50% шанс добавить красную строку
    if random.choice([True, False]):
        p.paragraph_format.first_line_indent = Cm(1.25) 
    else:
        p.paragraph_format.first_line_indent = Cm(0)  

    # разнообразие межстрочного интервала
    p.paragraph_format.space_after = Pt(random.randint(6, 12))

    # предотвращаем разрыв страницы внутри абзаца
    make_paragraph_keep_together(p)
    make_run_keep_with_next(run)

def add_footnote_to_cell(cell, text, footnote_number):
    """Добавляет сноску к элементу в ячейке таблицы."""
    p = cell.add_paragraph()
    run = p.add_run(f"[{footnote_number}]")
    run.font.superscript = True
    run.font.size = Pt(8)
    p.add_run(f" {text}")

def add_numbered_list(document, num_items=5):
    """Добавляет нумерованный список в документ."""
    for _ in range(num_items):
        sentence = fake.sentence(nb_words=6)[:50]  
        p = document.add_paragraph(sentence, style='List Number')
        
        # устанавливаем выравнивание и форматирование
        p.alignment = random.choice([WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY])
        p.runs[0].font.size = Pt(random.randint(10, 14))
        
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(0) 
        p.paragraph_format.line_spacing = Pt(12)  
        p.paragraph_format.space_before = Pt(0)  

    make_paragraph_keep_together(p)
    document.add_paragraph("")

def add_bulleted_list(document, num_items=5):
    """Добавляет маркированный список в документ."""
    for _ in range(num_items):
        sentence = fake.sentence(nb_words=6)[:50]  
        p = document.add_paragraph(sentence, style='List Bullet')

        p.alignment = random.choice([WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY])
        p.runs[0].font.size = Pt(random.randint(10, 14))

        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(0)  
        p.paragraph_format.line_spacing = Pt(12)  
        p.paragraph_format.space_before = Pt(0) 

    make_paragraph_keep_together(p)
    document.add_paragraph("")

def add_formula(document, generated_formulas):
    """Добавляет сгенерированную формулу как Office Math объект в документ."""
    formula = generate_complex_formula(generated_formulas)
    if not formula:
        return
    latex_formula = formula
    paragraph = document.add_paragraph()
    paragraph.space_after = Pt(random.randint(4, 12))
    paragraph.space_before = Pt(random.randint(4, 12))
    try:
        math2docx.add_math(paragraph, latex_formula)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception as e:
        pass

    # предотвращаем разрыв страницы между формулой и следующим элементом
    paragraph.paragraph_format.keep_with_next = True

def generate_complex_formula(generated_formulas):
    """Генерирует уникальную сложную LaTeX формулу."""
    formula_types = ['integral', 'sum', 'polynomial', 'fraction', 'product']
    formula_type = random.choice(formula_types)
    
    if formula_type == 'integral':
        a = random.randint(1, 10)
        b = random.randint(1, 20)
        c = random.randint(1, 10)
        formula = f"\\int_{{0}}^{{{b}}} {a}x^{{{c}}} \\, dx"
        
    elif formula_type == 'sum':
        n = random.randint(1, 20)
        power = random.randint(1, 5)
        formula = f"\\sum_{{i=1}}^{{{n}}} i^{{{power}}}"
        
    elif formula_type == 'polynomial':
        degree = random.randint(2, 5)
        coefficients = [random.randint(1, 10) for _ in range(degree + 1)]
        terms = []
        for i, coeff in enumerate(coefficients):
            power = degree - i
            if power > 1:
                terms.append(f"{coeff}x^{{{power}}}")
            elif power == 1:
                terms.append(f"{coeff}x")
            else:
                terms.append(f"{coeff}")
        formula = " + ".join(terms)
        
    elif formula_type == 'fraction':
        # генерация двух полиномов для числителя и знаменателя
        degree_num = random.randint(1, 3)
        degree_den = random.randint(1, 3)
        coefficients_num = [random.randint(1, 10) for _ in range(degree_num + 1)]
        coefficients_den = [random.randint(1, 10) for _ in range(degree_den + 1)]
        
        # создание числителя
        terms_num = []
        for i, coeff in enumerate(coefficients_num):
            power = degree_num - i
            if power > 1:
                terms_num.append(f"{coeff}x^{{{power}}}")
            elif power == 1:
                terms_num.append(f"{coeff}x")
            else:
                terms_num.append(f"{coeff}")
        numerator = " + ".join(terms_num)
        
        # создание знаменателя
        terms_den = []
        for i, coeff in enumerate(coefficients_den):
            power = degree_den - i
            if power > 1:
                terms_den.append(f"{coeff}x^{{{power}}}")
            elif power == 1:
                terms_den.append(f"{coeff}x")
            else:
                terms_den.append(f"{coeff}")
        denominator = " + ".join(terms_den)
        
        formula = f"\\frac{{{numerator}}}{{{denominator}}}"
        
    elif formula_type == 'product':
        # произведение двух полиномов
        degree_p1 = random.randint(1, 3)
        degree_p2 = random.randint(1, 3)
        coeff_p1 = random.randint(1, 10)
        coeff_p2 = random.randint(1, 10)
        
        # полином 1
        poly1_terms = []
        for i in range(degree_p1 +1):
            power = degree_p1 - i
            if power > 1:
                poly1_terms.append(f"{coeff_p1}x^{{{power}}}")
            elif power ==1:
                poly1_terms.append(f"{coeff_p1}x")
            else:
                poly1_terms.append(f"{coeff_p1}")
        poly1 = " + ".join(poly1_terms)
        
        # полином 2
        poly2_terms = []
        for i in range(degree_p2 +1):
            power = degree_p2 - i
            if power >1:
                poly2_terms.append(f"{coeff_p2}x^{{{power}}}")
            elif power ==1:
                poly2_terms.append(f"{coeff_p2}x")
            else:
                poly2_terms.append(f"{coeff_p2}")
        poly2 = " + ".join(poly2_terms)
        
        formula = f"({poly1}) \\cdot ({poly2})"
    else:
        return None

    # проверяем уникальность
    if formula in generated_formulas:
        return generate_complex_formula(generated_formulas)
    generated_formulas.add(formula)
    return formula

def generate_random_plot_image(doc_state):
    """Генерирует случайный график и сохраняет его во временный файл."""
    plot_type = random.choice(['line', 'bar', 'scatter', 'box', 'pie', 'hist'])
    fig, ax = plt.subplots(figsize=(4, 3))
    if plot_type == 'line':
        x = range(10)
        y = [random.randint(0, 100) for _ in x]
        ax.plot(x, y, marker='o')
        ax.set_title('Линейный график')
    elif plot_type == 'bar':
        categories = [fake.word() for _ in range(random.randint(3, 7))]
        values = [random.randint(10, 100) for _ in categories]
        ax.bar(categories, values, color=random.choice(['blue', 'green', 'red', 'cyan', 'magenta']))
        ax.set_title('Столбчатая диаграмма')
    elif plot_type == 'scatter':
        x = [random.uniform(0, 10) for _ in range(15)]
        y = [random.uniform(0, 10) for _ in range(15)]
        ax.scatter(x, y, color=random.choice(['blue', 'green', 'red', 'cyan', 'magenta']))
        ax.set_title('Точечная диаграмма')
    elif plot_type == 'box':
        data = [sorted([random.randint(0, 100) for _ in range(random.randint(10, 20))]) for _ in range(random.randint(2, 5))]
        ax.boxplot(data)
        ax.set_title('Боксплот')
    elif plot_type == 'pie':
        sizes = [random.randint(10, 100) for _ in range(random.randint(3, 6))]
        labels = [fake.word() for _ in sizes]
        ax.pie(sizes, labels=labels, autopct='%1.1f%%')
        ax.set_title('Круговая диаграмма')
    elif plot_type == 'hist':
        data = [random.gauss(50, 15) for _ in range(100)]
        ax.hist(data, bins=10, color=random.choice(['blue', 'green', 'red', 'cyan', 'magenta']))
        ax.set_title('Гистограмма')
    ax.axis('off')  # скрываем оси для чистоты изображения
    plt.tight_layout()

    # сохранение изображения с учётом режима колонок
    temp_dir = tempfile.gettempdir()
    unique_id = uuid.uuid4()
    image_path = os.path.join(temp_dir, f"plot_{unique_id}.png")
    plt.savefig(image_path, bbox_inches='tight', pad_inches=0.1)
    plt.close(fig)
    return image_path

def add_header_footer(document, header_text=None, footer_text=None):
    """Добавляет текст в колонтитулы документа с заданной вероятностью."""
    section = document.sections[0]
    
    # добавление header с вероятностью 50%
    if random.random() < 0.5 and header_text:
        header = section.header
        if not header.paragraphs:
            header_para = header.add_paragraph()
        else:
            header_para = header.paragraphs[0]
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_para.runs:
            header_para.runs[0].font.size = Pt(12)
        else:
            run = header_para.add_run(header_text)
            run.font.size = Pt(12)

    # добавление footer с вероятностью 70%
    if random.random() < 0.7 and footer_text:
        footer = section.footer
        if not footer.paragraphs:
            footer_para = footer.add_paragraph()
        else:
            footer_para = footer.paragraphs[0]
        footer_para.text = footer_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_para.runs:
            footer_para.runs[0].font.size = Pt(12)
        else:
            run = footer_para.add_run(footer_text)
            run.font.size = Pt(12)

def get_table_styles(document):
    """Получает все стили таблиц в документе."""
    styles = document.styles
    table_styles = []
    temp_table = document.add_table(rows=1, cols=1)
    for s in styles:
        if 'Table' in s.name or 'Таблица' in s.name:
            try:
                temp_table.style = s
                table_styles.append(s.name)
            except:
                pass
    # удаляем временную таблицу
    temp_table._element.getparent().remove(temp_table._element)
    return table_styles

def set_section_orientation(section, orientation=WD_ORIENT.PORTRAIT):
    """Устанавливает ориентацию секции документа."""
    new_width, new_height = (section.page_height, section.page_width) if orientation == WD_ORIENT.LANDSCAPE else (section.page_width, section.page_height)
    section.orientation = orientation
    section.page_width = new_width
    section.page_height = new_height

def estimate_lines_and_fullness(text, line_width=80):
    words = text.split()
    current_line_length = 0
    lines = []
    current_line = []
    
    for word in words:
        # +1 для пробела между словами
        if current_line_length + len(word) + (1 if current_line else 0) <= line_width:
            current_line.append(word)
            current_line_length += len(word) + (1 if current_line_length > 0 else 0)
        else:
            lines.append((current_line, current_line_length))
            current_line = [word]
            current_line_length = len(word)
    
    if current_line:
        lines.append((current_line, current_line_length))
    
    # подсчитываем количество полных строк (заполненность более 80%)
    full_lines = sum(1 for _, length in lines if length >= line_width * 0.8)
    
    return len(lines), full_lines

def is_suitable_for_justify(text, line_width=80):
    words = text.split()
    word_count = len(words)
    avg_word_length = sum(len(word) for word in words) / word_count if word_count > 0 else 0
    
    # базовые критерии
    min_words_per_line = 5
    min_avg_word_length = 4
    
    # проверяем количество строк и их заполненность
    total_lines, full_lines = estimate_lines_and_fullness(text, line_width)
    
    # примерное количество слов в строке
    words_per_line = word_count / total_lines if total_lines > 0 else 0
    
    return (words_per_line >= min_words_per_line and 
            avg_word_length >= min_avg_word_length and 
            word_count >= 15 and  # минимальное общее количество слов
            total_lines >= 3 and  # минимум 3 строки
            full_lines >= 2)      # минимум 2 полные строки

def add_random_elements(document, doc_state, end_footnotes, generated_formulas, last_element=None):
    """Добавляет случайные элементы в документ в зависимости от текущего состояния колонки."""
    try:
        # определяем доступные элементы
        if doc_state.in_multicol and any(count == 0 for count in doc_state.column_elements):
            elements = ['text', 'footnote']
            element_weights = {
                'text': 80,
                'footnote': 20
            }
        else:
            elements = ['heading', 'text', 'list', 'formula', 'footnote']
            if not doc_state.in_multicol:
                elements.extend(['image', 'table'])
            
            element_weights = {
                'heading': 10,
                'text': 65,
                'list': 10,
                'formula': 15,
                'image': 5,
                'table': 5,
                'footnote': 3
            }

        # исключаем элемент, если последний был таким же
        if last_element:
            elements = [elem for elem in elements if elem != last_element]
            weights = [element_weights[elem] for elem in elements]
        else:
            weights = [element_weights[elem] for elem in elements]

        element_type = random.choices(elements, weights=weights, k=1)[0]

        # добавляем выбранный элемент
        if element_type == 'heading' and not doc_state.in_multicol:
            try:
                level = random.choice([2, 3])
                heading_text = fake.sentence()
                p = document.add_heading(heading_text, level=level)
                if p:
                    p.paragraph_format.space_after = Pt(random.randint(4, 8))
                    p.paragraph_format.space_before = Pt(random.randint(9, 15))
                    p.paragraph_format.keep_with_next = True
            except Exception:
                print("Ошибка при добавлении заголовка")

        elif element_type == 'text':
            try:
                max_attempts = 10
                for _ in range(max_attempts):
                    text = fake.paragraph(nb_sentences=random.randint(3, 8))
                    suitable_for_justify = is_suitable_for_justify(text)
                    if suitable_for_justify or _ == max_attempts - 1:
                        break

                font_size = random.randint(10, 14)
                available_alignments = [
                    WD_ALIGN_PARAGRAPH.LEFT,
                    WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_PARAGRAPH.RIGHT
                ]
                if suitable_for_justify:
                    available_alignments.append(WD_ALIGN_PARAGRAPH.JUSTIFY)
                
                alignment = random.choice(available_alignments)

                if doc_state.in_multicol:
                    current_cell = doc_state.get_current_multicol_cell()
                    if current_cell:
                        p = add_text_to_cell(current_cell, text, font_size, alignment)
                        if p:
                            p.paragraph_format.space_after = Pt(random.randint(4, 8))
                else:
                    p = add_text(document, doc_state, text, font_size, alignment)
                    if p:
                        p.paragraph_format.space_after = Pt(random.randint(4, 8))
            except Exception:
                print("Ошибка при добавлении текста")

        elif element_type == 'list' and not doc_state.in_multicol:
            try:
                list_type = random.choice(['numbered', 'bulleted'])
                num_items = random.randint(3, 6)
                if list_type == 'numbered':
                    p = add_numbered_list(document, num_items)
                else:
                    p = add_bulleted_list(document, num_items)
                if p:
                    p.paragraph_format.space_after = Pt(random.randint(4, 8))
                    p.paragraph_format.space_before = Pt(random.randint(4, 8))
            except Exception:
                print("Ошибка при добавлении списка")

        elif element_type == 'formula' and not doc_state.in_multicol:
            try:
                p = add_formula(document, generated_formulas)
                if p:
                    p.paragraph_format.space_after = Pt(random.randint(4, 8))
                    p.paragraph_format.space_before = Pt(random.randint(4, 8))
            except Exception:
                print("Ошибка при добавлении формулы")

        elif element_type == 'footnote':
            try:
                footnote_text = fake.sentence(nb_words=10)
                if doc_state.in_multicol:
                    current_cell = doc_state.get_current_multicol_cell()
                    if current_cell:
                        p = add_footnote_to_cell(current_cell, footnote_text, len(end_footnotes) + 1)
                        if p:
                            p.paragraph_format.space_after = Pt(random.randint(4, 8))
                else:
                    p = document.add_paragraph()
                    if p:
                        p.paragraph_format.space_after = Pt(random.randint(4, 8))
                        add_footnote(p, footnote_text, len(end_footnotes) + 1)
                end_footnotes.append(footnote_text)
            except Exception:
                print("Ошибка при добавлении сноски")

        elif element_type == 'image' and not doc_state.in_multicol:
            try:
                p = add_image(document, doc_state)
                if p:
                    p.paragraph_format.space_after = Pt(random.randint(4, 8))
                    p.paragraph_format.space_before = Pt(random.randint(4, 8))
            except Exception:
                print("Ошибка при добавлении изображения")

        elif element_type == 'table' and not doc_state.in_multicol:
            try:
                p = add_table(document, doc_state)
                if p:
                    p.paragraph_format.space_after = Pt(random.randint(4, 8))
                    p.paragraph_format.space_before = Pt(random.randint(4, 8))
            except Exception:
                print("Ошибка при добавлении таблицы")

        return element_type

    except Exception as e:
        print(f"Общая ошибка в add_random_elements: {e}")
        return None

def generate_document(doc_id, output_dir, num_iterations=100):
    """Генерирует один документ с заданным количеством итераций добавления элементов."""
    try:
        document = Document()
        doc_state = DocumentState(document)

        # настройка базового стиля
        style = document.styles['Normal']
        style.paragraph_format.space_after = Pt(random.randint(4, 8))
        style.paragraph_format.space_before = Pt(random.randint(0, 2))

        # настройка стилей заголовков
        for i in range(1, 4):
            heading_style = document.styles[f'Heading {i}']
            heading_style.paragraph_format.space_after = Pt(random.randint(4, 8))
            heading_style.paragraph_format.space_before = Pt(random.randint(10, 14))
            heading_style.paragraph_format.keep_with_next = True

        # устанавливаем ориентацию первой секции
        set_section_orientation(document.sections[-1], WD_ORIENT.PORTRAIT)

        # добавление колонтитулов
        header_text = fake.sentence(nb_words=4)
        footer_text = fake.sentence(nb_words=4)
        add_header_footer(document, header_text=header_text, footer_text=footer_text)

        generated_formulas = set()
        end_footnotes = []
        last_element = None

        # начальный контент (одноколоночный)
        for _ in range(10):
            last_element = add_random_elements(document, doc_state, end_footnotes, generated_formulas, last_element)

        # основной цикл генерации контента
        for _ in range(num_iterations - 10):
            # отключаем разрыв страницы для текущего параграфа
            if document.paragraphs:
                current_paragraph = document.paragraphs[-1]
                if current_paragraph:
                    current_paragraph.paragraph_format.page_break_before = False

            if random.random() < 0.2:  # 20% шанс изменить ориентацию
                orientation = WD_ORIENT.LANDSCAPE if random.random() < 0.25 else WD_ORIENT.PORTRAIT
                section = document.add_section(WD_SECTION_START.CONTINUOUS)
                set_section_orientation(section, orientation)
                doc_state.current_orientation = orientation

            if doc_state.can_start_multicolumn() and not doc_state.in_multicol:
                if random.random() < 0.15:  # 15% шанс начать многоколонную секцию
                    possible_cols = []
                    if doc_state.can_add_multicolumn(2):
                        possible_cols.append(2)
                    if doc_state.can_add_multicolumn(3):
                        possible_cols.append(3)

                    if possible_cols:
                        num_cols = random.choice(possible_cols)
                        started = doc_state.start_multicolumn(num_cols)
                        
                        if started:
                            elements_per_col = doc_state.min_elements_per_column
                            total_elements = num_cols * elements_per_col

                            for _ in range(total_elements):
                                if random.random() < 0.8:  # 80% шанс для текста
                                    element_type = 'text'
                                else:  # 20% шанс для сноски
                                    element_type = 'footnote'

                                if element_type == 'text':
                                    text = fake.paragraph(nb_sentences=random.randint(3, 8))
                                    font_size = random.randint(10, 14)
                                    current_cell = doc_state.get_current_multicol_cell()
                                    if current_cell:
                                        p = add_text_to_cell(current_cell, text, font_size)
                                        if p:
                                            p.paragraph_format.space_after = Pt(random.randint(4, 8))
                                else:
                                    footnote_text = fake.sentence(nb_words=10)
                                    current_cell = doc_state.get_current_multicol_cell()
                                    if current_cell:
                                        p = add_footnote_to_cell(current_cell, footnote_text, len(end_footnotes) + 1)
                                        if p:
                                            p.paragraph_format.space_after = Pt(random.randint(4, 8))
                                        end_footnotes.append(footnote_text)

                                doc_state.elements_in_multicol += 1

                            doc_state.end_multicolumn()
                            continue

            # добавляем обычные элементы
            last_element = add_random_elements(document, doc_state, end_footnotes, generated_formulas, last_element)

        # проверка, чтобы документ не заканчивался многоколонной секцией
        if doc_state.in_multicol:
            doc_state.end_multicolumn()

        # добавляем концевые сноски
        if end_footnotes:
            if document.paragraphs:
                last_paragraph = document.paragraphs[-1]
                if last_paragraph:
                    last_paragraph.paragraph_format.page_break_before = False
            add_document_end_footnotes(document, end_footnotes)

        # cохранение документа
        output_path = os.path.join(output_dir, f'demo_{doc_id}.docx')
        document.save(output_path)
        return True

    except Exception as e:
        print(f"Ошибка при генерации документа {doc_id}: {e}")
        return False

def main():
    output_dir = './data/docx'
    os.makedirs(output_dir, exist_ok=True)
    num_docs = 5  # размер выборки, потом увеличим до 3000
    for doc_id in range(num_docs):
        # гарантируем, что многоколонные секции не будут на первой и последней страницах
        generate_document(doc_id, output_dir)
        if (doc_id + 1) % 10 == 0:
            print(f'Создано {doc_id + 1} документов')

if __name__ == "__main__":
    main()